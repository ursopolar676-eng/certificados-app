from flask import Flask, request, send_file, render_template_string
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os

app = Flask(__name__)

# carregar base
df = pd.read_excel("base.xlsx")

HTML = '''
<h2>Certificados</h2>
<form action="/gerar" method="post">
CPF:<br><input name="cpf"><br>
Data:<br><input name="data"><br>
Obra:<br><input name="obra"><br>
Endereço:<br><input name="endereco"><br>

Técnica:<br>
<select name="tecnica">
<option>kathlenn</option>
<option>vanusa</option>
<option>elaine</option>
<option>patricia</option>
<option>simone</option>
<option>marina</option>
</select><br><br>

<button type="submit">GERAR CERTIFICADO</button>
</form>
'''

@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/gerar", methods=["POST"])
def gerar():
    cpf = request.form["cpf"]
    data = request.form["data"]
    obra = request.form["obra"]
    endereco = request.form["endereco"]
    tecnica = request.form["tecnica"]

    pessoa = df[df["CPF"] == cpf]

    if pessoa.empty:
        return "CPF não encontrado"

    nome = pessoa.iloc[0]["NOME"]
    funcao = pessoa.iloc[0]["FUNCAO"]
    empresa = pessoa.iloc[0]["EMPRESA"]

    doc = Document("modelo.docx")

    for p in doc.paragraphs:
        p.text = p.text.replace("<<NOME>>", nome)
        p.text = p.text.replace("<<FUNCAO>>", funcao)
        p.text = p.text.replace("<<EMPRESA>>", empresa)
        p.text = p.text.replace("<<DATA>>", data)
        p.text = p.text.replace("<<OBRA>>", obra)
        p.text = p.text.replace("<<ENDERECO>>", endereco)

    assinatura = f"assinaturas/{tecnica}.png"

    # inserir assinatura no local correto
    for p in doc.paragraphs:
        if "<<ASSINATURA>>" in p.text:
            p.clear()
            run = p.add_run()
            run.add_picture(assinatura, width=Inches(2))

    docx_file = f"{nome}.docx"
    pdf_file = f"{nome}.pdf"

    doc.save(docx_file)
    convert(docx_file)

    return send_file(pdf_file, as_attachment=True)

app.run()
